import logging
import sys
import os
import uvicorn
import threading
from rembg import remove
from pathlib import Path
from docx import Document
from typing import Optional
from datetime import datetime
from pydantic import BaseModel
from PIL import ImageGrab, Image
from pynput import mouse, keyboard
from docx.shared import Inches, RGBColor
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, File, UploadFile, Form
from pymongo import MongoClient
from pydantic import BaseModel
import gridfs
from typing import Optional
import zipfile
from fastapi.responses import FileResponse
from io import BytesIO
from fastapi.responses import StreamingResponse
import shutil
from fastapi import APIRouter
from pydantic import BaseModel
import json

class UploadRequest(BaseModel):
    processId: str
    empId: str
    flowId: str

class RetrieveLogRequest(BaseModel):
    processId: str
    empId: str
    flowId: str


mongo_uri = "mongodb://localhost:27017/"
database_name = "ProcMap"

def get_gridfs_connection():
    client = MongoClient(mongo_uri)
    db = client[database_name]
    fs = gridfs.GridFS(db)
    return fs, db

doc = None
table = None
app = FastAPI()
cropped_screenshots = []
app_name = 'JARVIS - SMART'
mouse_listener_thread = None
Keyboard_listener_thread = None
# Flag to control the listener's state
listening = True
# Radius around the click to crop and typing delay
CROP_RADIUS = 50  # pixels
typing_delay = 2.0  # Seconds to wait before considering typing as stopped
# Buffer and lock for keystrokes, and timer
keystroke_buffer = []
buffer_lock = threading.Lock()
typing_timer = None
# A global variable to keep track of the listener thread
listener_thread: Optional[threading.Thread] = None

class LoggerStatus(BaseModel):
    is_active: bool

# Add CORSMiddleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

app_data_path = Path(os.getenv('APPDATA')) / app_name
full_screenshots_dir = app_data_path / 'Full_Screenshots'
cropped_screenshots_dir = app_data_path / 'Cropped_Screenshots'


# Before starting or stopping logging, adjust where the TagUI script file is created
activity_log_tag_path = app_data_path / 'Activity_Log.tag'
activity_log_tag2_path = app_data_path / 'Activity_Log2.tag'
activity_log_tag_path.parent.mkdir(parents=True, exist_ok=True)
activity_log_tag_path.parent.mkdir(parents=True, exist_ok=True)
activity_log_docx_path = app_data_path / 'Activity_Log.docx'
activity_log_tag2_content = ''



# log_directory = os.path.join(os.getenv('APPDATA', os.path.expanduser('~')), app_name)
# # Ensure the directory exists
# os.makedirs(log_directory, exist_ok=True)
# log_file_path = os.path.join(log_directory, 'app.log')
# Setup logging to the specified file
# logging.basicConfig(filename=log_file_path, filemode='a', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def append_to_tagui_script(content):
    with open(activity_log_tag_path, 'a') as file:
        file.write(content + "\n")

def remove_background(cropped_screenshot_path):
    cropped_screenshot_path = Path(cropped_screenshot_path)  # Convert to Path object if not already
    with open(cropped_screenshot_path, 'rb') as input_img:
        input_data = input_img.read()
    output_data = remove(input_data)
    # Generate output image path with new suffix
    output_image_path = cropped_screenshot_path.with_suffix('').with_name(cropped_screenshot_path.name + '-no-bg.png')
    with open(output_image_path, 'wb') as output_img:
        output_img.write(output_data)
    with open(activity_log_tag_path, 'a') as tagui_script_file:
        click_command = f'click {output_image_path}\n'
        tagui_script_file.write(click_command)

def add_row_to_table(description, full_path=None, cropped_path=None):
    row = table.add_row().cells
    row[0].text = description
    if full_path:
        paragraph = row[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(full_path, width=Inches(2.0))
        paragraph.add_run('\n' + os.path.basename(full_path))  # Add filename below the picture
    if cropped_path:
        paragraph = row[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(cropped_path, width=Inches(1.0))
        paragraph.add_run('\n' + os.path.basename(cropped_path))  # Add filename below the picture

def flush_keystrokes():
    global table, activity_log_tag2_content
    with buffer_lock:
        if keystroke_buffer:
            statement = ''.join(keystroke_buffer)
            # Create a new row for the typed statement
            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(f'Typed statement: "{statement}" on {datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.')
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set the color to red
            activity_log_tag2_content += f'keyboard {statement}\n'
            keystroke_buffer.clear()

def on_click(x, y, button, pressed):
    global activity_log_tag2_content
    if pressed and listening:
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        screenshot = ImageGrab.grab()
        full_screenshot_path = os.path.join(full_screenshots_dir, f'screenshot_{timestamp}.png')
        screenshot.save(full_screenshot_path)
        # Cropping logic
        left, top, right, bottom = max(x - CROP_RADIUS, 0), max(y - CROP_RADIUS, 0), x + CROP_RADIUS, y + CROP_RADIUS
        cropped_screenshot = screenshot.crop((left, top, right, bottom))
        cropped_screenshot_path = os.path.join(cropped_screenshots_dir, f'cropped_{timestamp}.png')
        cropped_screenshot.save(cropped_screenshot_path)
        # Add the click action and screenshots to the table
        action_description = f'Click at ({x}, {y}) with {button} on {timestamp}.'
        add_row_to_table(action_description, full_screenshot_path, cropped_screenshot_path)
        if cropped_screenshot_path:  # Ensure path exists
            cropped_screenshots.append(cropped_screenshot_path)
        activity_log_tag2_content += f'click {cropped_screenshot_path}\nwait 4\n'    

def on_scroll(x, y, dx, dy):
    if listening:
        # Log scroll action
        scroll_direction = 'down' if dy < 0 else 'up'
        description = f'Scroll {scroll_direction} at ({x}, {y}) on {datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.'
        add_row_to_table(description)

def on_press(key):
    global typing_timer
    spKeys = False
    if listening:
        with buffer_lock:
            try:
                # If backspace is pressed, remove the last character from buffer
                if key == keyboard.Key.backspace:
                    if keystroke_buffer:  # Check if buffer is not empty
                        keystroke_buffer.pop()  # Remove the last character
                elif key == keyboard.Key.enter:
                    spKeys = True
                    keystroke_buffer.append('[enter]\nwait 3\n')
                elif key == keyboard.Key.tab:
                    spKeys = True
                    keystroke_buffer.append('[tab]\nwait 3\n')
                elif key == keyboard.Key.space:
                    keystroke_buffer.append(' ')
                elif key == keyboard.Key.cmd_l or key == keyboard.Key.cmd_r:
                    spKeys = True
                    keystroke_buffer.append('[win]')
                elif key == keyboard.Key.ctrl:
                    keystroke_buffer.append('[ctrl]')
                elif key == keyboard.Key.alt:
                    keystroke_buffer.append('[alt]')
                elif key == keyboard.Key.esc:
                    spKeys = True
                    keystroke_buffer.append('[esc]')
                else:
                    keystroke_buffer.append(key.char)
            except AttributeError:
                if key == keyboard.Key.space:
                    keystroke_buffer.append(' ')  # Add space for readability

        # Restart the typing timer
        if typing_timer:
            typing_timer.cancel()
        if not spKeys:    
            typing_timer = threading.Timer(typing_delay, flush_keystrokes)
            typing_timer.start()
        else:
            flush_keystrokes()    


def stop_listening(key):#
    global listening
    if key == keyboard.Key.esc:
        listening = False
        if typing_timer:
            typing_timer.cancel()
            flush_keystrokes()  # Ensure we flush any remaining keystrokes
        return False


@app.post("/upload/")
async def upload_file(empId: str = Form(...), processId: str = Form(...), flowId: str = Form(...), file: UploadFile = File(...)):

    # Now you directly use the form fields to construct your UploadRequest object
    # upload_request = UploadRequest(empId=empId, processId=processId, flowId=flowId)

    fs, db = get_gridfs_connection()
    employees_collection = db["employees"]

    # The rest of your function remains unchanged, using the constructed UploadRequest object
    employee = employees_collection.find_one({"empId": empId})
    if not employee:
        raise HTTPException(status_code=404, detail="Employee not found.")

    processObject = next((p for p in employee.get("processObjects", []) if p["processId"] == processId), None)
    if not processObject:
        raise HTTPException(status_code=404, detail="Process not found.")

    # Save uploaded file to GridFS
    contents = await file.read()
    file_id = fs.put(contents, filename=file.filename)

    # Update the database entry
    processObject.setdefault('flows', []).append({"file_id": file_id, "flowId": flowId})
    employees_collection.update_one({"empId": empId}, {"$set": {"processObjects": employee["processObjects"]}})

    return {"message": "File uploaded successfully", "file_id": str(file_id)}

@app.get("/retrieve/")
def retrieve_file(upload_request: RetrieveLogRequest):
    fs, db = get_gridfs_connection()
    employees_collection = db["employees"]

    employee = employees_collection.find_one({"empId": upload_request.empId})
    if not employee:
        raise HTTPException(status_code=404, detail="Employee not found.")

    processObject = next((p for p in employee.get("processObjects", []) if p["processId"] == upload_request.processId), None)
    # if not processObject or 'file_id' not in processObject:
    #     raise HTTPException(status_code=404, detail="Process or file not found.")
    if processObject:
        flow = next((f for f in processObject.get("flows",[]) if f["flowId"] == upload_request.flowId),None)
        if not flow or 'file_id' not in flow:
            raise HTTPException(status_code=404, detail="Process or file not found.")

        file_id = flow['file_id']
        file = fs.get(file_id)

        return StreamingResponse(BytesIO(file.read()), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=download.zip"})
    else:
        raise HTTPException(status_code=404, detail="Process not found.")


@app.post("/start-logging/")
async def start_logging():
    global activity_log_tag2_content,listener_thread, listening, doc, table, cropped_screenshots_dir, full_screenshots_dir, activity_log_docx_path, activity_log_tag_path, cropped_screenshots, Keyboard_listener_thread, mouse_listener_thread

    # Cleanup before starting
    if activity_log_docx_path.exists():
        activity_log_docx_path.unlink()
    if activity_log_tag_path.exists():
        activity_log_tag_path.unlink()
    if activity_log_tag2_path.exists():
        activity_log_tag2_path.unlink()    

    shutil.rmtree(full_screenshots_dir, ignore_errors=True)
    shutil.rmtree(cropped_screenshots_dir, ignore_errors=True)
    full_screenshots_dir.mkdir(parents=True, exist_ok=True)
    cropped_screenshots_dir.mkdir(parents=True, exist_ok=True)

    # Initialize doc and table
    doc = Document()
    doc.add_heading('Activity Log', 0)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Action Description'
    hdr_cells[1].text = 'Full Screenshot'
    hdr_cells[2].text = 'Cropped Screenshot'
    cropped_screenshots = []
    activity_log_tag2_content = ''

    #delete all the threads
    # if listener_thread:
    #     listener_thread.join()
    listener_thread = None

    # if mouse_listener_thread:
    #     mouse_listener_thread.join()
    mouse_listener_thread = None

    # if Keyboard_listener_thread:
    #     Keyboard_listener_thread.join()
    Keyboard_listener_thread = None


    if not listener_thread or not listener_thread.is_alive():
        listening = True
        listener_thread = threading.Thread(target=run_listeners)
        listener_thread.start()
        return {"message": "Logging started."}
    else:
        raise HTTPException(status_code=400, detail="Logging is already in progress.")

@app.post("/stop-logging/")
async def stop_logging():
    global mouse_listener, keyboard_listener

    if mouse_listener is not None:
        mouse_listener.stop()
        mouse_listener.join()  # Wait for the listener thread to fully stop
        mouse_listener = None

    if keyboard_listener is not None:
        keyboard_listener.stop()
        keyboard_listener.join()
        keyboard_listener = None

    # for screenshot in cropped_screenshots:
    #     remove_background(screenshot)
    if doc:
        doc.save(activity_log_docx_path)

    with open(activity_log_tag_path, 'w') as file:
        file.write(activity_log_tag2_content)    

    return {"message": "Logging stopped and data saved."}

def run_listeners():
    global mouse_listener, keyboard_listener
    mouse_listener = mouse.Listener(on_click=on_click, on_scroll=on_scroll)
    mouse_listener.start()

    keyboard_listener = keyboard.Listener(on_press=on_press, on_release=stop_listening)
    keyboard_listener.start()


if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8002)

