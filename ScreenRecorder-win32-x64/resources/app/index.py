from fastapi import FastAPI, HTTPException
import uvicorn
import threading
from pydantic import BaseModel
from typing import Optional
import os
import io
import time
import threading
from docx import Document
from PIL import ImageGrab, Image
from datetime import datetime
from pynput import mouse, keyboard
from docx.shared import Inches, RGBColor
from rembg import remove
import shutil


# Assuming other necessary imports and your defined functions are above

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Add CORSMiddleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)


def remove_background(cropped_screenshot_path):
    # Open the input image
    with open(cropped_screenshot_path, 'rb') as input_img:
        input_data = input_img.read()
    # Use rembg to remove the background
    output_data = remove(input_data)
    # Save the output image
    output_image_path = cropped_screenshot_path+'-no-bg.png'
    with open(output_image_path, 'wb') as output_img:
        output_img.write(output_data)
    click_command = f'click {output_image_path}\n'
    tagui_script_file.write(click_command)


mouse_listener_thread = None
Keyboard_listener_thread = None

# Initialize a Word document and add a heading
cropped_screenshots = []
# doc.add_heading('Activity Log', 0)

doc = None
table = None

from pathlib import Path

app_name = 'JARVIS - SMART'
app_data_path = Path(os.getenv('APPDATA')) / app_name
full_screenshots_dir = app_data_path / 'Full_Screenshots'
cropped_screenshots_dir = app_data_path / 'Cropped_Screenshots'

# Create the directories, ensuring parent directories are created too
full_screenshots_dir.mkdir(parents=True, exist_ok=True)
cropped_screenshots_dir.mkdir(parents=True, exist_ok=True)


# Flag to control the listener's state
listening = True

# Radius around the click to crop and typing delay
CROP_RADIUS = 50  # pixels
typing_delay = 2.0  # Seconds to wait before considering typing as stopped

# Buffer and lock for keystrokes, and timer
keystroke_buffer = []
buffer_lock = threading.Lock()
typing_timer = None

def add_row_to_table(description, full_path=None, cropped_path=None):
    row = table.add_row().cells
    row[0].text = description
    if full_path:
        paragraph = row[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(full_path, width=Inches(2.0))
        paragraph.add_run('\n' + os.path.basename(full_path))  # Add filename below the picture
    if cropped_path:
        paragraph = row[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(cropped_path, width=Inches(1.0))
        paragraph.add_run('\n' + os.path.basename(cropped_path))  # Add filename below the picture


def flush_keystrokes():
    global table
    with buffer_lock:
        if keystroke_buffer:
            statement = ''.join(keystroke_buffer)
            # Create a new row for the typed statement
            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            run = paragraph.add_run(f'Typed statement: "{statement}" on {datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.')
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set the color to red
            keystroke_buffer.clear()

def on_click(x, y, button, pressed):
    if pressed and listening:
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        screenshot = ImageGrab.grab()
        full_screenshot_path = os.path.join(full_screenshots_dir, f'screenshot_{timestamp}.png')
        screenshot.save(full_screenshot_path)
        # Cropping logic
        left, top, right, bottom = max(x - CROP_RADIUS, 0), max(y - CROP_RADIUS, 0), x + CROP_RADIUS, y + CROP_RADIUS
        cropped_screenshot = screenshot.crop((left, top, right, bottom))
        cropped_screenshot_path = os.path.join(cropped_screenshots_dir, f'cropped_{timestamp}.png')
        cropped_screenshot.save(cropped_screenshot_path)
        # Add the click action and screenshots to the table
        action_description = f'Click at ({x}, {y}) with {button} on {timestamp}.'
        add_row_to_table(action_description, full_screenshot_path, cropped_screenshot_path)
        if cropped_screenshot_path:  # Ensure path exists
            cropped_screenshots.append(cropped_screenshot_path)

def on_scroll(x, y, dx, dy):
    if listening:
        # Log scroll action
        scroll_direction = 'down' if dy < 0 else 'up'
        description = f'Scroll {scroll_direction} at ({x}, {y}) on {datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.'
        add_row_to_table(description)

def on_press(key):
    global typing_timer
    if listening:
        with buffer_lock:
            try:
                # If backspace is pressed, remove the last character from buffer
                if key == keyboard.Key.backspace:
                    if keystroke_buffer:  # Check if buffer is not empty
                        keystroke_buffer.pop()  # Remove the last character
                else:
                    keystroke_buffer.append(key.char)
            except AttributeError:
                if key == keyboard.Key.space:
                    keystroke_buffer.append(' ')  # Add space for readability

        # Restart the typing timer
        if typing_timer:
            typing_timer.cancel()
        typing_timer = threading.Timer(typing_delay, flush_keystrokes)
        typing_timer.start()


def stop_listening(key):
    global listening
    if key == keyboard.Key.esc:
        listening = False
        if typing_timer:
            typing_timer.cancel()
            flush_keystrokes()  # Ensure we flush any remaining keystrokes
        return False

# A global variable to keep track of the listener thread
listener_thread: Optional[threading.Thread] = None

class LoggerStatus(BaseModel):
    is_active: bool

@app.post("/start-logging/")
async def start_logging():
    global listener_thread, listening, tagui_script_file, doc, table, cropped_screenshots

    # delete Full_Screenshots and Cropped_Screenshots directories if they exist and create new ones
    if os.path.exists(full_screenshots_dir):
        shutil.rmtree(full_screenshots_dir)
    os.makedirs(full_screenshots_dir)
    if os.path.exists(cropped_screenshots_dir):
        shutil.rmtree(cropped_screenshots_dir)
    os.makedirs(cropped_screenshots_dir)

    #delete the Activity_Log.docx and Activity_Log.tag file if it exists
    if os.path.exists('Activity_Log.docx'):
        os.remove('Activity_Log.docx')
    if os.path.exists('Activity_Log.tag'):
        os.remove('Activity_Log.tag')

    doc = None
    table = None
    cropped_screenshots = []

    doc = Document()
    doc.add_heading('Activity Log', 0)

    # Add a table with a header row
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Action Description'
    hdr_cells[1].text = 'Full Screenshot'
    hdr_cells[2].text = 'Cropped Screenshot'



    if listener_thread is None or not listener_thread.is_alive():
        listening = True
        # Initialize your document and table as before
        doc = Document()
        doc.add_heading('Activity Log', 0)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Action Description'
        hdr_cells[1].text = 'Full Screenshot'
        hdr_cells[2].text = 'Cropped Screenshot'

        tagui_script_path = 'Activity_Log.tag'
        tagui_script_file = open(tagui_script_path, 'w')

        listener_thread = threading.Thread(target=run_listeners)
        listener_thread.start()
        return {"message": "Logging started."}
    else:
        raise HTTPException(status_code=400, detail="Logging is already in progress.")

@app.post("/stop-logging/")
async def stop_logging():
    global listening, mouse_listener_thread, Keyboard_listener_thread, listener_thread, tagui_script_file, doc

    # try:
    listening = False

    # Wait for the listener thread to stop
    if mouse_listener_thread:
        mouse_listener_thread.join()
    if Keyboard_listener_thread:
        Keyboard_listener_thread.join()
    listener_thread.join()
    # Once the thread stops, you might want to save your document and do cleanup
    for screenshot in cropped_screenshots:
        remove_background(screenshot)
    tagui_script_file.close()
    print("TagUI script file closed.")
    if doc:
        doc.save('Activity_Log.docx')
    if tagui_script_file:
        tagui_script_file.close()
    # if os.path.exists(full_screenshots_dir):
    #     shutil.rmtree(full_screenshots_dir)
    # if os.path.exists(cropped_screenshots_dir):
    #     shutil.rmtree(cropped_screenshots_dir)

    return {"message": "Logging stopped and data saved."}
    # Except Exception as e:

def run_listeners():
    # You need to modify this part to initialize your listeners
    # For example:
    global listening, mouse_listener_thread, Keyboard_listener_thread
    try:
        mouse_listener_thread = threading.Thread(target=lambda: mouse.Listener(on_click=on_click, on_scroll=on_scroll).start())
        mouse_listener_thread.start()

        Keyboard_listener_thread = threading.Thread(target=lambda: keyboard.Listener(on_press=on_press, on_release=stop_listening).start())
        Keyboard_listener_thread.start()
    finally:
        # Cleanup actions, such as saving the document and closing files
        pass

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8002)
